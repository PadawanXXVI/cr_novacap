# ===========================================
# MÓDULO: gerar_relatorio_sei.py
# ===========================================
# Gera relatórios institucionais padronizados SEI-GDF (.docx)
# Emitido por: NOVACAP/PRES/CPCR
# Destinado a: Diretorias (DC, DO, DP, DS etc.)
# ===========================================

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# -------------------------------------------------------------
# Dicionário de responsabilidades por Diretoria
# -------------------------------------------------------------
MAPEAMENTO_DIRETORIAS = {
    "Alambrado (Cercamento)": "DC",
    "Doação de Mudas": "DC",
    "Jardim": "DC",
    "Mato Alto": "DC",
    "Meio-fio": "DC",
    "Parque Infantil": "DC",
    "Pista de Skate": "DC",
    "Poda / Supressão de Árvore": "DC",
    "Ponto de Encontro Comunitário (PEC)": "DC",
    "Praça": "DC",
    "Quadra de Esporte": "DC",
    "Tapa-buraco": "DC",

    "Boca de Lobo": "DO",
    "Bueiro": "DO",
    "Calçada": "DO",
    "Estacionamentos": "DO",
    "Galeria de Águas Pluviais": "DO",
    "Passagem Subterrânea": "DO",
    "Passarela": "DO",
    "Pisos Articulados": "DO",
    "Rampa": "DO",
    "Rua, Via ou Rodovia (Pista)": "DO",
    "Limpeza de Resíduos da Novacap": "DO",
    "Ciclovia ou Ciclofaixa (pista)": "DO",
}


# -------------------------------------------------------------
# Função principal
# -------------------------------------------------------------
def gerar_relatorio_sei(df: pd.DataFrame, filtros: dict, autor: str) -> str:
    """
    Gera um relatório Word (.docx) no padrão SEI-GDF.
    Emitido pela NOVACAP/PRES/CPCR com base em filtros aplicados.
    
    Parâmetros:
        df (pd.DataFrame): DataFrame com os dados filtrados
        filtros (dict): dicionário com status, RA, diretoria, datas etc.
        autor (str): nome de usuário que gerou o relatório
    Retorna:
        Caminho do arquivo gerado (.docx)
    """

    # ---------------------------------------------------------
    # 1️⃣ Início do documento
    # ---------------------------------------------------------
    doc = Document()

    # Define margens SEI
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(1.18)

    # ---------------------------------------------------------
    # 2️⃣ Cabeçalho institucional
    # ---------------------------------------------------------
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("Governo do Distrito Federal\n")
    run.bold = True
    run.font.size = Pt(12)

    run = titulo.add_run("Companhia Urbanizadora da Nova Capital do Brasil – NOVACAP\n")
    run.bold = True
    run.font.size = Pt(12)

    run = titulo.add_run("Comissão Permanente da Central de Relacionamento – CPCR")
    run.font.size = Pt(11)

    doc.add_paragraph("\n")

    # ---------------------------------------------------------
    # 3️⃣ Identificação e assunto
    # ---------------------------------------------------------
    numero_relatorio = datetime.now().strftime("%m/%Y")
    p_titulo = doc.add_paragraph()
    p_titulo.add_run(f"Relatório nº {numero_relatorio} – NOVACAP/PRES/CPCR").bold = True

    data_br = datetime.now().strftime("%d de %B de %Y").capitalize()
    p_data = doc.add_paragraph(f"Brasília, {data_br}.")
    p_data.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    assunto = "Consolidação das Demandas"
    if filtros.get("status"):
        assunto += f" – Status: {filtros['status']}"
    if filtros.get("ra"):
        assunto += f" – RA: {filtros['ra']}"
    if filtros.get("diretoria"):
        assunto += f" – Diretoria: {filtros['diretoria']}"

    doc.add_paragraph(f"\nAssunto: {assunto}")
    doc.add_paragraph("Ilmo. Sr. Diretor-Presidente da NOVACAP,\n")

    # ---------------------------------------------------------
    # 4️⃣ SEÇÃO 1: CONTEXTO
    # ---------------------------------------------------------
    doc.add_paragraph("1. CONTEXTO").runs[0].bold = True
    p1 = doc.add_paragraph(
        "1.1. O presente relatório, elaborado pela Comissão Permanente da Central de Relacionamento (CPCR/PRES), "
        "tem por finalidade consolidar as informações referentes às solicitações registradas e tramitadas "
        "no Sistema de Controle de Processos e Protocolo de Atendimentos – SISTRAMITE, "
        "encaminhadas às Diretorias competentes da NOVACAP, conforme os filtros aplicados neste relatório."
    )
    p1.paragraph_format.space_after = Pt(8)

    # ---------------------------------------------------------
    # 5️⃣ SEÇÃO 2: RELATO
    # ---------------------------------------------------------
    doc.add_paragraph("2. RELATO").runs[0].bold = True

    if df.empty:
        doc.add_paragraph(
            "2.1. Não foram encontrados registros para os parâmetros selecionados, "
            "razão pela qual este relatório se limita à formalização da ausência de movimentações no período."
        )
    else:
        total = len(df)
        por_status = df["Status"].value_counts().to_dict()
        principais_status = ", ".join([f"{k}: {v}" for k, v in por_status.items()])

        doc.add_paragraph(
            f"2.1. Foram processadas {total} solicitações conforme os filtros definidos. "
            f"A distribuição por status é a seguinte: {principais_status}."
        )

        # Totais por RA (se houver)
        if "RA" in df.columns:
            por_ra = df["RA"].value_counts().to_dict()
            top_ra = ", ".join([f"{k} ({v})" for k, v in list(por_ra.items())[:5]])
            doc.add_paragraph(
                f"2.2. As Regiões Administrativas com maior número de solicitações são: {top_ra}."
            )

        # Totais por responsável (opcional)
        if "Responsável" in df.columns:
            por_resp = df["Responsável"].value_counts().to_dict()
            doc.add_paragraph(
                f"2.3. A distribuição por responsáveis técnicos indica {len(por_resp)} servidores distintos atuantes."
            )

        # -----------------------------------------------------
        # Agrupamento por Diretoria (automático via mapeamento)
        # -----------------------------------------------------
        if "Serviço" in df.columns:
            df["Diretoria"] = df["Serviço"].map(MAPEAMENTO_DIRETORIAS)
            grupos = df.groupby("Diretoria")

            for diretoria, grupo in grupos:
                if pd.isna(diretoria):
                    continue

                doc.add_paragraph(f"\n2.{diretoria} – Demandas sob responsabilidade da Diretoria {diretoria}").runs[0].bold = True
                tabela = doc.add_table(rows=1, cols=5)
                tabela.style = "Table Grid"
                hdr = tabela.rows[0].cells
                hdr[0].text = "Serviço"
                hdr[1].text = "RA"
                hdr[2].text = "Status"
                hdr[3].text = "Responsável"
                hdr[4].text = "Data"

                for _, linha in grupo.iterrows():
                    row = tabela.add_row().cells
                    row[0].text = str(linha.get("Serviço", ""))
                    row[1].text = str(linha.get("RA", ""))
                    row[2].text = str(linha.get("Status", ""))
                    row[3].text = str(linha.get("Responsável", ""))
                    row[4].text = str(linha.get("Data", ""))

                doc.add_paragraph("\n")

    # ---------------------------------------------------------
    # 6️⃣ SEÇÃO 3: CONCLUSÃO
    # ---------------------------------------------------------
    doc.add_paragraph("3. CONCLUSÃO").runs[0].bold = True
    doc.add_paragraph(
        "3.1. Recomenda-se o acompanhamento periódico das solicitações, priorizando as demandas de "
        "execução direta pelas Diretorias das Cidades (DC), Obras (DO), Planejamento (DP) e Suporte (DS), "
        "bem como o reforço na comunicação entre a CPCR e as Regiões Administrativas, "
        "com vistas à melhoria contínua do fluxo de atendimento e à redução de devoluções."
    )

    # ---------------------------------------------------------
    # 7️⃣ Rodapé institucional
    # ---------------------------------------------------------
    doc.add_paragraph("\nAtenciosamente,\n\n")

    rodape = (
        "\"Brasília - Patrimônio Cultural da Humanidade\"\n"
        "Setor de Áreas Públicas - Lote B - Bairro Guará - CEP 70075-900 - DF\n"
        "Telefone(s): (61) XXXX-XXXX\n"
        "Sítio - www.novacap.df.gov.br"
    )
    doc.add_paragraph(rodape)
    linha = doc.add_paragraph("__________________________________________________________")
    linha.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Doc. SEI/GDF XXXXXXXX").alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Metadados finais
    criacao = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
    doc.add_paragraph(f"Criado por {autor}, versão 1 em {criacao}.")

    # ---------------------------------------------------------
    # 8️⃣ Salvamento
    # ---------------------------------------------------------
    nome_arquivo = f"Relatorio_SEI_CPCR_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(nome_arquivo)

    return nome_arquivo
